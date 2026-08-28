#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Gerador Fiel do Ementário do PPC em Formato Word (.docx):
Recorte literal da Seção de Ementas da Estrutura Curricular do PPC Técnico em Administração (IFSC Garopaba).
Lê diretamente de 'ementario_adm.tex' e produz o documento Word com a exata estrutura e formatação do PPC.
"""

import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

BASE_DIR = "/Users/chameoandre/Google-Drive-chameoandre/INSTITUTO-FEDERAL-SANTA-CATARINA/CURSOS/TECNICO/informatica-integrado/reformulacao-ppc-informatica-administracao-integrado/tecnico-administracao"
TEX_PATH = os.path.join(BASE_DIR, "documento-ppc-principal", "ementario_adm.tex")
CHECK_DIR = os.path.join(BASE_DIR, "revisao-equipe", "checagem-biblioteca")
DOCX_OUTPUT_PATH = os.path.join(CHECK_DIR, "Ementario_Completo_PPC_Tecnico_Administracao_Revisao_Biblioteca.docx")

# Colors matching IFSC PPC
COLOR_IFSC_GREEN = RGBColor(16, 140, 80)      # #108C50
COLOR_DARK = RGBColor(15, 23, 42)             # #0F172A
COLOR_MUTED = RGBColor(100, 116, 139)         # #64748B
HEX_IFSC_GREEN = "108C50"
HEX_LIGHT_GREEN = "EBF8F0"
HEX_LIGHT_GRAY = "F1F5F9"
HEX_BORDER = "CBD5E1"

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def clean_tex(text):
    if not text: return ""
    text = re.sub(r'\\revisao\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\textcolor\{[^}]*\}\{([^}]*)\}', r'\1', text)
    text = text.replace(r'\textbf{', '**').replace('}', '**')
    text = text.replace(r'\textit{', '*').replace('}', '*')
    text = text.replace(r'\large', '').replace(r'\small', '').replace(r'\normalsize', '')
    text = text.replace(r'\newline', '\n').replace(r'\\[3pt]', '\n').replace(r'\\', '\n')
    text = text.replace(r'\vspace{2pt}', '').replace(r'\vspace{0.2cm}', '')
    text = text.replace(r'\&', '&').replace(r'\%', '%').replace(r'\$', '$').replace(r'\_', '_')
    text = text.replace('--', '—')
    return text.strip()

def add_formatted_text(paragraph, text, default_size=9.5, default_color=COLOR_DARK):
    # Splits by **bold**
    tokens = re.split(r'(\*\*[^*]+\*\*)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.font.bold = True
        else:
            # check for *italic*
            it_tokens = re.split(r'(\*[^*]+\*)', token)
            for it_t in it_tokens:
                if it_t.startswith('*') and it_t.endswith('*'):
                    run = paragraph.add_run(it_t[1:-1])
                    run.font.italic = True
                else:
                    run = paragraph.add_run(it_t)
        if len(paragraph.runs) > 0:
            paragraph.runs[-1].font.size = Pt(default_size)
            if default_color:
                paragraph.runs[-1].font.color.rgb = default_color

def parse_tex_tables():
    with open(TEX_PATH, "r", encoding="utf-8") as f:
        content = f.read()

    # Split ementas tables
    parts = content.split(r'\begin{xltabular}{\linewidth}{|X|p{2.5cm}|p{2.5cm}|}')
    intro_part = parts[0]
    
    uc_tables = []
    for tab in parts[1:]:
        tab_clean = tab.split(r'\end{xltabular}')[0].strip()
        
        # 1. Parse header: UC, Semestre, CH EaD, CH Total
        uc_match = re.search(r'\\textbf\{\\large\s*([^}]+)\}', tab_clean)
        if not uc_match:
            uc_match = re.search(r'\\textbf\{([^}]+)\}\s*\}\s*&\s*\\multicolumn', tab_clean)
        uc_nome = uc_match.group(1).strip() if uc_match else "Unidade Curricular"
        uc_nome = clean_tex(uc_nome).replace('**', '').strip()
        
        sem_match = re.search(r'\\textbf\{Semestre:\}\}\s*\\textbf\{([^}]+)\}', tab_clean)
        semestre = sem_match.group(1).strip() if sem_match else ""
        
        ead_match = re.search(r'CH EaD\*:\}\}\s*\\\\\s*\\cline\{2-3\}\s*&\s*\\textbf\{([^}]+)\}', tab_clean)
        ch_ead = ead_match.group(1).strip() if ead_match else "00 h"
        
        tot_match = re.search(r'CH Total\*:\}\}\s*\\\\\s*\\cline\{2-3\}\s*&\s*[^\&]+\&\s*\\textbf\{([^}]+)\}', tab_clean)
        ch_total = tot_match.group(1).strip() if tot_match else ""
        
        # Extract sections: Objetivos, Conteúdos, Estratégias, Básica, Complementar
        def extract_sec(title, next_title=None):
            pattern = rf'\\textbf\{{{title}\:\}}\}}\s*\\\\\s*\\hline\s*\\multicolumn\{{3\}}\{{\|p\{{[^}}]+\}}\|\}}{{\s*(.*?)\s*}}\s*\\\\\s*\\hline'
            m = re.search(pattern, tab_clean, re.DOTALL)
            if m:
                return m.group(1).strip()
            return ""

        raw_obj = extract_sec('Objetivos')
        # Parse itemize items
        obj_items = []
        if r'\begin{itemize}' in raw_obj:
            items_raw = raw_obj.split(r'\item')
            for it in items_raw[1:]:
                it_c = clean_tex(it.replace(r'\end{itemize}', '')).replace('\n', ' ').strip()
                if it_c: obj_items.append(it_c)
        else:
            obj_items = [clean_tex(raw_obj)]

        raw_cont = extract_sec('Conteúdos')
        cont_clean = clean_tex(raw_cont)

        raw_estr = extract_sec('Estratégias de Ensino e Aprendizagem')
        estr_clean = clean_tex(raw_estr)

        raw_bb = extract_sec('Bibliografia Básica')
        bb_items = [clean_tex(b).strip() for b in raw_bb.split(r'\newline') if clean_tex(b).strip()]

        raw_bc = extract_sec('Bibliografia Complementar')
        bc_items = [clean_tex(c).strip() for c in raw_bc.split(r'\newline') if clean_tex(c).strip()]

        uc_tables.append({
            "uc_nome": uc_nome,
            "semestre": semestre,
            "ch_ead": ch_ead,
            "ch_total": ch_total,
            "objetivos": obj_items,
            "conteudos": cont_clean,
            "estrategias": estr_clean,
            "basica": bb_items,
            "complementar": bc_items
        })
        
    return intro_part, uc_tables

def generate_literal_ppc_docx():
    print(f"Gerando Recorte Literal do Ementário do PPC em: {DOCX_OUTPUT_PATH}")
    intro_part, uc_tables = parse_tex_tables()
    
    doc = Document()
    
    # Set Standard A4 Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(9.5)
    style_normal.font.color.rgb = COLOR_DARK

    # Cabeçalho Oficial do PPC
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_h1 = p_head.add_run("INSTITUTO FEDERAL DE SANTA CATARINA — CÂMPUS GAROPABA\n")
    r_h1.font.bold = True
    r_h1.font.size = Pt(11)
    r_h1.font.color.rgb = COLOR_IFSC_GREEN
    
    r_h2 = p_head.add_run("PROJETO PEDAGÓGICO DE CURSO (PPC 2026)\nTÉCNICO INTEGRADO EM ADMINISTRAÇÃO\n")
    r_h2.font.bold = True
    r_h2.font.size = Pt(13)
    r_h2.font.color.rgb = COLOR_DARK
    
    r_h3 = p_head.add_run("Seção da Estrutura Curricular & Ementário das 45 Unidades Curriculares\n")
    r_h3.font.bold = True
    r_h3.font.size = Pt(10.5)
    r_h3.font.color.rgb = COLOR_MUTED

    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(12)
    
    # 1. Matrizes Curriculares Anuais
    p_sec1 = doc.add_paragraph()
    r_s1 = p_sec1.add_run("Matriz Curricular Detalhada por Anos e Blocos:")
    r_s1.font.bold = True
    r_s1.font.size = Pt(11.5)
    r_s1.font.color.rgb = COLOR_IFSC_GREEN
    p_sec1.paragraph_format.space_after = Pt(8)

    # Função auxiliar para tabela de matriz curricular
    def add_matriz_table(ano_titulo, rows_data, subtotal_text, subtotal_val):
        tbl = doc.add_table(rows=len(rows_data) + 3, cols=4)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl, color=HEX_BORDER)
        
        # Header Row 1: Banner
        cell_top = tbl.cell(0, 0)
        cell_top.merge(tbl.cell(0, 3))
        set_cell_background(cell_top, HEX_LIGHT_GREEN)
        set_cell_margins(cell_top, top=80, bottom=80, left=100, right=100)
        p = cell_top.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(ano_titulo)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_IFSC_GREEN
        
        # Header Row 2: Columns
        col_headers = ["Ano", "Unidade Curricular (UC)", "Bloco Curricular", "CH (h)"]
        col_widths = [Inches(1.1), Inches(3.2), Inches(1.8), Inches(0.9)]
        for c_i in range(4):
            c_h = tbl.cell(1, c_i)
            c_h.width = col_widths[c_i]
            set_cell_background(c_h, HEX_LIGHT_GRAY)
            set_cell_margins(c_h, top=60, bottom=60, left=80, right=80)
            p = c_h.paragraphs[0]
            r = p.add_run(col_headers[c_i])
            r.font.bold = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = COLOR_MUTED
            
        for r_i, row in enumerate(rows_data, 2):
            for c_i in range(4):
                c = tbl.cell(r_i, c_i)
                c.width = col_widths[c_i]
                set_cell_margins(c, top=45, bottom=45, left=80, right=80)
                p = c.paragraphs[0]
                r = p.add_run(row[c_i])
                r.font.size = Pt(8.5)
                if c_i in (0, 3):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
        # Subtotal row
        r_sub = len(rows_data) + 2
        c_sub_label = tbl.cell(r_sub, 0)
        c_sub_label.merge(tbl.cell(r_sub, 2))
        set_cell_background(c_sub_label, HEX_LIGHT_GRAY)
        set_cell_margins(c_sub_label, top=60, bottom=60, left=80, right=80)
        p_sl = c_sub_label.paragraphs[0]
        p_sl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_sl = p_sl.add_run(subtotal_text)
        r_sl.font.bold = True
        r_sl.font.size = Pt(8.5)
        
        c_sub_val = tbl.cell(r_sub, 3)
        set_cell_background(c_sub_val, HEX_LIGHT_GRAY)
        set_cell_margins(c_sub_val, top=60, bottom=60, left=80, right=80)
        p_sv = c_sub_val.paragraphs[0]
        p_sv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sv = p_sv.add_run(subtotal_val)
        r_sv.font.bold = True
        r_sv.font.size = Pt(8.5)

    # 1º Ano
    matriz_1 = [
        ["1º Ano", "Artes — Ano 1", "Formação Geral", "80 h"],
        ["1º Ano", "Educação Física — Ano 1", "Formação Geral", "80 h"],
        ["1º Ano", "Inglês — Ano 1", "Formação Geral", "80 h"],
        ["1º Ano", "Língua Portuguesa e Literatura — Ano 1", "Formação Geral", "80 h"],
        ["1º Ano", "Espanhol — Ano 1", "Formação Geral", "80 h"],
        ["1º Ano", "Biologia — Ano 1", "Formação Geral", "80 h"],
        ["1º Ano", "Física — Ano 1", "Formação Geral", "40 h"],
        ["1º Ano", "Matemática — Ano 1", "Formação Geral", "80 h"],
        ["1º Ano", "Química — Ano 1", "Formação Geral", "40 h"],
        ["1º Ano", "Filosofia — Ano 1", "Formação Geral", "80 h"],
        ["1º Ano", "Geografia — Ano 1", "Formação Geral", "80 h"],
        ["1º Ano", "Sociologia — Ano 1", "Formação Geral", "80 h"],
        ["1º Ano", "Introdução à Administração", "Formação Técnica", "80 h"],
        ["1º Ano", "Sociedade e Trabalho", "Formação Técnica", "40 h"],
        ["1º Ano", "Gestão de Marketing I", "Formação Técnica", "40 h"],
        ["1º Ano", "Organização e Processos", "Formação Técnica", "40 h"],
        ["1º Ano", "Informática Aplicada", "Formação Técnica", "40 h"]
    ]
    add_matriz_table("1º ANO LETIVO (SEMESTRES 1 E 2)", matriz_1, "Subtotal 1º Ano (880h FG + 240h FT):", "1.120 h")
    
    p_sp1 = doc.add_paragraph()
    p_sp1.paragraph_format.space_before = Pt(10)

    # 2º Ano
    matriz_2 = [
        ["2º Ano", "Artes — Ano 2", "Formação Geral", "80 h"],
        ["2º Ano", "Educação Física — Ano 2", "Formação Geral", "80 h"],
        ["2º Ano", "Inglês — Ano 2", "Formação Geral", "80 h"],
        ["2º Ano", "Língua Portuguesa e Literatura — Ano 2", "Formação Geral", "80 h"],
        ["2º Ano", "Espanhol — Ano 2", "Formação Geral", "80 h"],
        ["2º Ano", "Física — Ano 2", "Formação Geral", "80 h"],
        ["2º Ano", "Matemática — Ano 2", "Formação Geral", "80 h"],
        ["2º Ano", "Química — Ano 2", "Formação Geral", "80 h"],
        ["2º Ano", "História — Ano 2", "Formação Geral", "80 h"],
        ["2º Ano", "Sociologia — Ano 2", "Formação Geral", "80 h"],
        ["2º Ano", "Matemática para Administração", "Formação Técnica", "40 h"],
        ["2º Ano", "Gestão de Marketing II", "Formação Técnica", "40 h"],
        ["2º Ano", "Gestão de Operações e Qualidade", "Formação Técnica", "80 h"],
        ["2º Ano", "Empreendedorismo I", "Formação Técnica", "40 h"],
        ["2º Ano", "Responsabilidade Socioambiental e Sustentabilidade", "Formação Técnica", "40 h"],
        ["2º Ano", "Oficina de Integração I", "Núcleo Politécnico", "80 h"]
    ]
    add_matriz_table("2º ANO LETIVO (SEMESTRES 3 E 4)", matriz_2, "Subtotal 2º Ano (800h FG + 240h FT + 80h NP):", "1.120 h")

    p_sp2 = doc.add_paragraph()
    p_sp2.paragraph_format.space_before = Pt(10)

    # 3º Ano
    matriz_3 = [
        ["3º Ano", "Língua Portuguesa e Literatura — Ano 3", "Formação Geral", "80 h"],
        ["3º Ano", "Biologia — Ano 3", "Formação Geral", "80 h"],
        ["3º Ano", "Física — Ano 3", "Formação Geral", "80 h"],
        ["3º Ano", "Matemática — Ano 3", "Formação Geral", "80 h"],
        ["3º Ano", "Química — Ano 3", "Formação Geral", "40 h"],
        ["3º Ano", "Filosofia — Ano 3", "Formação Geral", "80 h"],
        ["3º Ano", "Geografia — Ano 3", "Formação Geral", "80 h"],
        ["3º Ano", "História — Ano 3", "Formação Geral", "80 h"],
        ["3º Ano", "Empreendedorismo II", "Formação Técnica", "40 h"],
        ["3º Ano", "Gestão de Pessoas e Relações no Trabalho", "Formação Técnica", "80 h"],
        ["3º Ano", "Gestão Financeira", "Formação Técnica", "80 h"],
        ["3º Ano", "Oficina de Integração II", "Núcleo Politécnico", "80 h"]
    ]
    add_matriz_table("3º ANO LETIVO (SEMESTRES 5 E 6)", matriz_3, "Subtotal 3º Ano (600h FG + 200h FT + 80h NP):", "960 h")

    doc.add_page_break()

    # 2. Seção do Ementário das 45 Unidades Curriculares
    p_sec2 = doc.add_paragraph()
    r_s2 = p_sec2.add_run("Ementário das Unidades Curriculares (Formato Visual Oficial em Tabela)")
    r_s2.font.bold = True
    r_s2.font.size = Pt(12.5)
    r_s2.font.color.rgb = COLOR_IFSC_GREEN
    p_sec2.paragraph_format.space_after = Pt(12)

    # Render each of the 45 UC tables exactly as defined in the PPC LaTeX format
    for idx, uc in enumerate(uc_tables, 1):
        # Table with 3 columns (matching |X|p{2.5cm}|p{2.5cm}|)
        tbl_uc = doc.add_table(rows=8, cols=3)
        tbl_uc.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl_uc, color=HEX_BORDER)
        
        w_col0 = Inches(4.2)
        w_col1 = Inches(1.3)
        w_col2 = Inches(1.3)
        
        # Row 0: Header top (Unidade Curricular merged vertically in TeX; in Word we structure clean cells)
        c_uc = tbl_uc.cell(0, 0)
        c_uc.width = w_col0
        set_cell_background(c_uc, HEX_LIGHT_GREEN)
        set_cell_margins(c_uc, top=70, bottom=70, left=90, right=90)
        p_u = c_uc.paragraphs[0]
        r_u_tag = p_u.add_run("Unidade Curricular:\n")
        r_u_tag.font.bold = True
        r_u_tag.font.size = Pt(8.5)
        r_u_tag.font.color.rgb = COLOR_IFSC_GREEN
        r_u_name = p_u.add_run(uc["uc_nome"])
        r_u_name.font.bold = True
        r_u_name.font.size = Pt(11)
        r_u_name.font.color.rgb = COLOR_DARK
        
        c_sem = tbl_uc.cell(0, 1)
        c_sem.merge(tbl_uc.cell(0, 2))
        set_cell_margins(c_sem, top=70, bottom=70, left=90, right=90)
        p_s = c_sem.paragraphs[0]
        r_s_tag = p_s.add_run("Semestre: ")
        r_s_tag.font.bold = True
        r_s_tag.font.size = Pt(8.5)
        r_s_tag.font.color.rgb = COLOR_IFSC_GREEN
        r_s_val = p_s.add_run(uc["semestre"])
        r_s_val.font.bold = True
        r_s_val.font.size = Pt(9.5)

        # Row 1: CH EaD e CH Total
        c_r1_0 = tbl_uc.cell(1, 0)
        set_cell_margins(c_r1_0, top=40, bottom=40, left=90, right=90)
        p_r1_0 = c_r1_0.paragraphs[0]
        p_r1_0.add_run("Componente Curricular da Matriz Oficial").font.size = Pt(8)
        p_r1_0.runs[0].font.color.rgb = COLOR_MUTED
        
        c_ead = tbl_uc.cell(1, 1)
        set_cell_margins(c_ead, top=40, bottom=40, left=90, right=90)
        p_ead = c_ead.paragraphs[0]
        p_ead.add_run("CH EaD*: ").font.bold = True
        p_ead.runs[0].font.color.rgb = COLOR_IFSC_GREEN
        p_ead.runs[0].font.size = Pt(8.5)
        p_ead.add_run(uc["ch_ead"]).font.bold = True
        p_ead.runs[1].font.size = Pt(9)
        
        c_tot = tbl_uc.cell(1, 2)
        set_cell_margins(c_tot, top=40, bottom=40, left=90, right=90)
        p_tot = c_tot.paragraphs[0]
        p_tot.add_run("CH Total*: ").font.bold = True
        p_tot.runs[0].font.color.rgb = COLOR_IFSC_GREEN
        p_tot.runs[0].font.size = Pt(8.5)
        p_tot.add_run(uc["ch_total"]).font.bold = True
        p_tot.runs[1].font.size = Pt(9)

        # Row 2: Objetivos
        c_obj = tbl_uc.cell(2, 0)
        c_obj.merge(tbl_uc.cell(2, 2))
        set_cell_margins(c_obj, top=70, bottom=70, left=90, right=90)
        p_obj_h = c_obj.paragraphs[0]
        r_oh = p_obj_h.add_run("Objetivos:\n")
        r_oh.font.bold = True
        r_oh.font.size = Pt(9)
        r_oh.font.color.rgb = COLOR_IFSC_GREEN
        
        for obj_it in uc["objetivos"]:
            p_it = c_obj.add_paragraph()
            p_it.paragraph_format.left_indent = Inches(0.2)
            p_it.paragraph_format.space_after = Pt(2)
            p_it.add_run("• ")
            add_formatted_text(p_it, obj_it, default_size=9)

        # Row 3: Conteúdos
        c_cont = tbl_uc.cell(3, 0)
        c_cont.merge(tbl_uc.cell(3, 2))
        set_cell_margins(c_cont, top=70, bottom=70, left=90, right=90)
        p_cont_h = c_cont.paragraphs[0]
        r_ch = p_cont_h.add_run("Conteúdos:\n")
        r_ch.font.bold = True
        r_ch.font.size = Pt(9)
        r_ch.font.color.rgb = COLOR_IFSC_GREEN
        
        for cont_line in uc["conteudos"].split('\n'):
            if cont_line.strip():
                p_c = c_cont.add_paragraph()
                p_c.paragraph_format.space_after = Pt(3)
                add_formatted_text(p_c, cont_line.strip(), default_size=9)

        # Row 4: Estratégias de Ensino e Aprendizagem
        c_estr = tbl_uc.cell(4, 0)
        c_estr.merge(tbl_uc.cell(4, 2))
        set_cell_margins(c_estr, top=70, bottom=70, left=90, right=90)
        p_estr_h = c_estr.paragraphs[0]
        r_eh = p_estr_h.add_run("Estratégias de Ensino e Aprendizagem:\n")
        r_eh.font.bold = True
        r_eh.font.size = Pt(9)
        r_eh.font.color.rgb = COLOR_IFSC_GREEN
        
        for estr_line in uc["estrategias"].split('\n'):
            if estr_line.strip():
                p_e = c_estr.add_paragraph()
                p_e.paragraph_format.space_after = Pt(3)
                add_formatted_text(p_e, estr_line.strip(), default_size=9)

        # Row 5: Bibliografia Básica
        c_bb = tbl_uc.cell(5, 0)
        c_bb.merge(tbl_uc.cell(5, 2))
        set_cell_margins(c_bb, top=70, bottom=70, left=90, right=90)
        set_cell_background(c_bb, HEX_LIGHT_GREEN)
        p_bb_h = c_bb.paragraphs[0]
        r_bbh = p_bb_h.add_run("Bibliografia Básica:\n")
        r_bbh.font.bold = True
        r_bbh.font.size = Pt(9)
        r_bbh.font.color.rgb = COLOR_IFSC_GREEN
        
        for bb_item in uc["basica"]:
            if bb_item.strip():
                p_b = c_bb.add_paragraph()
                p_b.paragraph_format.space_after = Pt(4)
                add_formatted_text(p_b, bb_item.strip(), default_size=9)

        # Row 6: Bibliografia Complementar
        c_bc = tbl_uc.cell(6, 0)
        c_bc.merge(tbl_uc.cell(6, 2))
        set_cell_margins(c_bc, top=70, bottom=70, left=90, right=90)
        p_bc_h = c_bc.paragraphs[0]
        r_bch = p_bc_h.add_run("Bibliografia Complementar:\n")
        r_bch.font.bold = True
        r_bch.font.size = Pt(9)
        r_bch.font.color.rgb = COLOR_IFSC_GREEN
        
        for bc_item in uc["complementar"]:
            if bc_item.strip():
                p_c = c_bc.add_paragraph()
                p_c.paragraph_format.space_after = Pt(4)
                add_formatted_text(p_c, bc_item.strip(), default_size=9)
                
        # Row 7: Empty divider row (or remove it)
        c_r7 = tbl_uc.cell(7, 0)
        c_r7.merge(tbl_uc.cell(7, 2))
        set_cell_margins(c_r7, top=20, bottom=20, left=90, right=90)
        set_cell_background(c_r7, HEX_LIGHT_GRAY)
        p_end = c_r7.paragraphs[0]
        p_end.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_end.add_run(f"PPC Técnico em Administração (IFSC Garopaba) • UC {idx:02d}").font.size = Pt(7.5)
        p_end.runs[0].font.color.rgb = COLOR_MUTED

        # Page break after each UC (except the last one)
        if idx < len(uc_tables):
            doc.add_page_break()

    doc.save(DOCX_OUTPUT_PATH)
    print(f"Documento Word (.docx) gerado com sucesso em:\n{DOCX_OUTPUT_PATH}")

if __name__ == "__main__":
    generate_literal_ppc_docx()
