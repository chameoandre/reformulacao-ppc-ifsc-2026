#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Gerador do Relatório Oficial de Adequação e Validação de Bibliografias para o NDE
Curso Superior de Tecnologia em Gestão Ambiental (PPC 2026) - IFSC Câmpus Garopaba.

Baseado fielmente no modelo institucional:
'modelos-de-documento/Relatório de Adequação de Bibliografias.docx'

Gera documento Word (.docx) padronizado para as 33 Unidades Curriculares, com cruzamento
automático dos quantitativos do acervo físico (Sophia) e biblioteca virtual.
100% Otimizado e compatível com Google Docs, Microsoft Word e LibreOffice (bordas e cores preservadas).
"""

import os
import re
import unicodedata
import pandas as pd
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

CHECK_DIR = os.path.dirname(os.path.abspath(__file__))
TXT_PATH = os.path.join(CHECK_DIR, "Unidades Curriculares Gestão Ambiental.txt")
EXCEL_PATH = os.path.join(CHECK_DIR, "Analise_Bibliografica_PPC_vs_Acervo_Sophia_GA.xlsx")
OUTPUT_DOCX_PATH = os.path.join(CHECK_DIR, "Relatorio_Adequacao_Bibliografias_NDE_CST_Gestao_Ambiental.docx")

# Cores e Estilos
HEX_GRAY_HEADER = "D9D9D9"
HEX_BORDER = "000000"
COLOR_DARK = RGBColor(15, 23, 42)
COLOR_IFSC_GREEN = RGBColor(16, 140, 80)

NDE_MEMBERS = [
    "1. Eduardo Cargnin Ferreira (NDE)",
    "2. Elisa Serena Gandolfo Martins (NDE)",
    "3. Fabiana de Agapito Kangerski (NDE)",
    "4. Jaciara Zarpellon Mazo (NDE)",
    "5. Jean Marcel de Almeida Espinoza (NDE)",
    "6. Juliano da Cunha Gomes (NDE)",
    "7. Nauber Gavski da Silva (NDE)",
    "8. Renato da Silva Rosa Rodrigues (NDE)"
]

def normalize_str(text):
    if not text:
        return ""
    text = unicodedata.normalize("NFKD", str(text)).encode("ASCII", "ignore").decode("ASCII").lower()
    text = re.sub(r"[^a-z0-9\s]", " ", text)
    return re.sub(r"\s+", " ", text).strip()

def set_cell_background(cell, hex_color="D9D9D9"):
    """Define o preenchimento da célula compatível com o Google Docs e MS Word"""
    tcPr = cell._tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')):
        tcPr.remove(s)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    """Define margens internas da célula"""
    tcPr = cell._tc.get_or_add_tcPr()
    for m in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(m)
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, color="000000", sz="6", val="single"):
    """
    Garante que cada célula possua bordas explícitas (<w:tcBorders>),
    evitando que o Google Docs remova ou ignore as linhas da tabela.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    for b in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(b)
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def set_table_borders(table, color="000000", sz="6", val="single"):
    """Configura o estilo TableGrid e as bordas gerais da tabela"""
    tblPr = table._tbl.tblPr
    
    # tblStyle TableGrid
    style_el = tblPr.find(qn('w:tblStyle'))
    if style_el is None:
        style_el = parse_xml(f'<w:tblStyle {nsdecls("w")} w:val="TableGrid"/>')
        tblPr.insert(0, style_el)
    else:
        style_el.set(qn('w:val'), 'TableGrid')
        
    # tblLayout fixed
    layout_el = tblPr.find(qn('w:tblLayout'))
    if layout_el is None:
        layout_el = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
        tblPr.append(layout_el)

    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'</w:tblBorders>'
    )
    for b in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(b)
    tblPr.append(borders)

try:
    import analise_acervo_biblioteca_ga as ga
except ImportError:
    ga = None

def get_bold_title_runs(ref_raw):
    raw = ref_raw.strip()
    if any(k in raw for k in ["LIVRO didático", "Material didático", "Fundo Nacional"]):
        return [(raw, False)]
        
    meta = ga.extract_metadata(raw) if ga else {}
    titulo = meta.get("titulo", "").strip()
    titulo_curto = meta.get("titulo_curto", "").strip()
    
    candidates = [titulo, titulo_curto]
    for cand in candidates:
        if not cand: continue
        idx = raw.find(cand)
        if idx != -1:
            prefix = raw[:idx]
            title_text = raw[idx:idx+len(cand)]
            suffix = raw[idx+len(cand):]
            runs = []
            if prefix: runs.append((prefix, False))
            if title_text: runs.append((title_text, True))
            if suffix: runs.append((suffix, False))
            return runs
            
        pattern = re.escape(cand)
        m = re.search(pattern, raw, re.IGNORECASE)
        if m:
            prefix = raw[:m.start()]
            title_text = raw[m.start():m.end()]
            suffix = raw[m.end():]
            runs = []
            if prefix: runs.append((prefix, False))
            if title_text: runs.append((title_text, True))
            if suffix: runs.append((suffix, False))
            return runs
            
    m = re.match(r'^([A-ZÁÉÍÓÚÂÊÔÃÕÇ\s\.,;&\(\)\-]+?\.\s+)(.+?)(?=\.\s+(?:[A-Z][a-z]+|\d+\.\s*ed|[A-ZÁÉÍÓÚÂÊÔÃÕÇ\s]+:|$))(.*)$', raw)
    if m:
        return [(m.group(1), False), (m.group(2), True), (m.group(3), False)]
        
    return [(raw, False)]

def add_abnt_reference_paragraph(cell, ref_num, ref_text):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    
    # Prefix [1]
    r_num = p.add_run(f"[{ref_num}] ")
    r_num.font.name = 'Arial'
    r_num.font.size = Pt(9.5)
    r_num.font.bold = True
    
    raw = ref_text.strip()
    runs_data = get_bold_title_runs(raw)
    for text_part, is_bold in runs_data:
        r = p.add_run(text_part)
        r.font.name = 'Arial'
        r.font.size = Pt(9.5)
        if is_bold:
            r.font.bold = True
    return p

def load_all_data():
    from analise_acervo_biblioteca_ga import parse_txt_ementas
    ucs_info, _ = parse_txt_ementas()
    df_all = pd.read_excel(EXCEL_PATH, sheet_name='Diagnóstico Consolidado')
    return ucs_info, df_all

def determine_availability_and_type(row_data, ref_text):
    """
    Retorna (Quantidade_Disponivel_Str, Tipo_Físico_Virtual_Str)
    com base no diagnóstico da auditoria do acervo e metadados da obra.
    """
    exemplares = row_data.get('Exemplares_Acervo', 0)
    if pd.isna(exemplares):
        exemplares = 0
    else:
        try:
            exemplares = int(exemplares)
        except:
            exemplares = 0
            
    status = str(row_data.get('Status', ''))
    
    ref_lower = ref_text.lower()
    is_digital_ref = any(k in ref_lower for k in [
        "http://", "https://", "disponível em:", "acesso em:", "scielo", 
        "ibge", "mma.gov", "planalto.gov", "leis.gov", "periódicos", "revista"
    ])
    
    if exemplares > 0:
        qtd_str = str(exemplares)
        tipo_str = "F / V" if is_digital_ref else "F"
    elif "PNLD" in ref_text or "LIVRO didático" in ref_text:
        qtd_str = "PNLD / FNDE"
        tipo_str = "F"
    elif is_digital_ref or "online" in status.lower():
        qtd_str = "Virtual / Digital"
        tipo_str = "V"
    else:
        qtd_str = "0 (A adquirir)"
        tipo_str = "F"
        
    return qtd_str, tipo_str

def build_uc_nde_table(doc, uc, df_uc):
    """
    Constrói a tabela padronizada do NDE para uma UC específica.
    """
    # Linha de data
    p_data = doc.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_data.paragraph_format.space_before = Pt(4)
    p_data.paragraph_format.space_after = Pt(2)
    r_dt = p_data.add_run("Data: \t___/09/2026.")
    r_dt.font.name = 'Arial'
    r_dt.font.size = Pt(10)
    
    # 4 colunas base
    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=HEX_BORDER, sz="6", val="single")
    
    # ROW 0: Cabeçalho Principal (span 4)
    row0 = table.add_row()
    c0 = row0.cells[0].merge(row0.cells[1]).merge(row0.cells[2]).merge(row0.cells[3])
    set_cell_background(c0, HEX_GRAY_HEADER)
    set_cell_margins(c0, top=100, bottom=100, left=120, right=120)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(2)
    p0.paragraph_format.space_after = Pt(2)
    r0 = p0.add_run("Análise da Adequação e Validação das Bibliografias")
    r0.font.name = 'Arial'
    r0.font.size = Pt(11.5)
    r0.font.bold = True
    
    # ROW 1: Unidade Curricular | CH | Semestre
    row1 = table.add_row()
    c_uc = row1.cells[0].merge(row1.cells[1])
    c_ch = row1.cells[2]
    c_sem = row1.cells[3]
    
    set_cell_background(c_uc, HEX_GRAY_HEADER)
    set_cell_background(c_ch, HEX_GRAY_HEADER)
    set_cell_background(c_sem, HEX_GRAY_HEADER)
    
    set_cell_margins(c_uc, top=80, bottom=80, left=100, right=100)
    set_cell_margins(c_ch, top=80, bottom=80, left=80, right=80)
    set_cell_margins(c_sem, top=80, bottom=80, left=80, right=80)
    
    p_uc = c_uc.paragraphs[0]
    p_uc.paragraph_format.space_after = Pt(0)
    r_ucl = p_uc.add_run("Unidade Curricular: ")
    r_ucl.font.name = 'Arial'
    r_ucl.font.size = Pt(10)
    r_ucl.font.bold = True
    r_ucv = p_uc.add_run(uc['uc_nome'])
    r_ucv.font.name = 'Arial'
    r_ucv.font.size = Pt(10)
    r_ucv.font.bold = True
    
    p_ch = c_ch.paragraphs[0]
    p_ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ch.paragraph_format.space_after = Pt(0)
    r_chl = p_ch.add_run("CH: ")
    r_chl.font.name = 'Arial'
    r_chl.font.size = Pt(10)
    r_chl.font.bold = True
    r_chv = p_ch.add_run(f"{uc['ch_total']}")
    r_chv.font.name = 'Arial'
    r_chv.font.size = Pt(10)
    
    p_sem = c_sem.paragraphs[0]
    p_sem.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sem.paragraph_format.space_after = Pt(0)
    r_seml = p_sem.add_run("Semestre: ")
    r_seml.font.name = 'Arial'
    r_seml.font.size = Pt(10)
    r_seml.font.bold = True
    r_semv = p_sem.add_run(f"{uc['semestre']}º")
    r_semv.font.name = 'Arial'
    r_semv.font.size = Pt(10)
    
    # ROW 2: Conteúdos (span 4)
    row2 = table.add_row()
    c_cont = row2.cells[0].merge(row2.cells[1]).merge(row2.cells[2]).merge(row2.cells[3])
    set_cell_margins(c_cont, top=100, bottom=100, left=120, right=120)
    p_cont = c_cont.paragraphs[0]
    p_cont.paragraph_format.space_after = Pt(2)
    p_cont.paragraph_format.line_spacing = 1.15
    r_ctl = p_cont.add_run("Conteúdos:\n")
    r_ctl.font.name = 'Arial'
    r_ctl.font.size = Pt(10)
    r_ctl.font.bold = True
    
    r_ctv = p_cont.add_run(uc['conteudos'] if uc['conteudos'] else "Conforme matriz curricular aprovada.")
    r_ctv.font.name = 'Arial'
    r_ctv.font.size = Pt(9.5)
    
    # ROW 3: Cabeçalho Bibliografia Básica
    row3 = table.add_row()
    c_bb_head = row3.cells[0].merge(row3.cells[1])
    c_bb_qtd = row3.cells[2]
    c_bb_fv = row3.cells[3]
    
    set_cell_background(c_bb_head, HEX_GRAY_HEADER)
    set_cell_background(c_bb_qtd, HEX_GRAY_HEADER)
    set_cell_background(c_bb_fv, HEX_GRAY_HEADER)
    
    set_cell_margins(c_bb_head, top=80, bottom=80, left=100, right=100)
    set_cell_margins(c_bb_qtd, top=80, bottom=80, left=60, right=60)
    set_cell_margins(c_bb_fv, top=80, bottom=80, left=60, right=60)
    
    p_bbh = c_bb_head.paragraphs[0]
    p_bbh.paragraph_format.space_after = Pt(0)
    r_bbh = p_bbh.add_run("Referências Bibliografia Básica")
    r_bbh.font.name = 'Arial'
    r_bbh.font.size = Pt(10)
    r_bbh.font.bold = True
    
    p_bbq = c_bb_qtd.paragraphs[0]
    p_bbq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bbq.paragraph_format.space_after = Pt(0)
    r_bbq = p_bbq.add_run("Quantidade Disponível")
    r_bbq.font.name = 'Arial'
    r_bbq.font.size = Pt(9)
    r_bbq.font.bold = True
    
    p_bbf = c_bb_fv.paragraphs[0]
    p_bbf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bbf.paragraph_format.space_after = Pt(0)
    r_bbf = p_bbf.add_run("Físico (F)\nVirtual (V)")
    r_bbf.font.name = 'Arial'
    r_bbf.font.size = Pt(8.5)
    r_bbf.font.bold = True
    
    # Rows BB
    df_bb = df_uc[df_uc['Tipo'] == 'BÁSICA'].reset_index(drop=True)
    ref_idx = 1
    for i, bb_ref in enumerate(uc['bb']):
        row_bb = table.add_row()
        c_ref = row_bb.cells[0].merge(row_bb.cells[1])
        c_qtd = row_bb.cells[2]
        c_fv = row_bb.cells[3]
        
        set_cell_margins(c_ref, top=60, bottom=60, left=80, right=80)
        set_cell_margins(c_qtd, top=60, bottom=60, left=40, right=40)
        set_cell_margins(c_fv, top=60, bottom=60, left=40, right=40)
        
        row_data = df_bb.iloc[i].to_dict() if i < len(df_bb) else {}
        qtd_str, tipo_str = determine_availability_and_type(row_data, bb_ref)
        
        add_abnt_reference_paragraph(c_ref, ref_idx, bb_ref)
        
        p_q = c_qtd.paragraphs[0]
        p_q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_q.paragraph_format.space_after = Pt(0)
        r_q = p_q.add_run(qtd_str)
        r_q.font.name = 'Arial'
        r_q.font.size = Pt(9.5)
        if "0" in qtd_str:
            r_q.font.color.rgb = RGBColor(185, 28, 28)
            r_q.font.bold = True
        
        p_f = c_fv.paragraphs[0]
        p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_f.paragraph_format.space_after = Pt(0)
        r_f = p_f.add_run(tipo_str)
        r_f.font.name = 'Arial'
        r_f.font.size = Pt(9.5)
        r_f.font.bold = True
        
        ref_idx += 1

    # ROW Cabeçalho Bibliografia Complementar
    row_bc_head = table.add_row()
    c_bc_head = row_bc_head.cells[0].merge(row_bc_head.cells[1])
    c_bc_qtd = row_bc_head.cells[2]
    c_bc_fv = row_bc_head.cells[3]
    
    set_cell_background(c_bc_head, HEX_GRAY_HEADER)
    set_cell_background(c_bc_qtd, HEX_GRAY_HEADER)
    set_cell_background(c_bc_fv, HEX_GRAY_HEADER)
    
    set_cell_margins(c_bc_head, top=80, bottom=80, left=100, right=100)
    set_cell_margins(c_bc_qtd, top=80, bottom=80, left=60, right=60)
    set_cell_margins(c_bc_fv, top=80, bottom=80, left=60, right=60)
    
    p_bch = c_bc_head.paragraphs[0]
    p_bch.paragraph_format.space_after = Pt(0)
    r_bch = p_bch.add_run("Referências Bibliografia Complementar")
    r_bch.font.name = 'Arial'
    r_bch.font.size = Pt(10)
    r_bch.font.bold = True
    
    p_bcq = c_bc_qtd.paragraphs[0]
    p_bcq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bcq.paragraph_format.space_after = Pt(0)
    r_bcq = p_bcq.add_run("Quantidade Disponível")
    r_bcq.font.name = 'Arial'
    r_bcq.font.size = Pt(9)
    r_bcq.font.bold = True
    
    p_bcf = c_bc_fv.paragraphs[0]
    p_bcf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bcf.paragraph_format.space_after = Pt(0)
    r_bcf = p_bcf.add_run("Físico (F)\nVirtual (V)")
    r_bcf.font.name = 'Arial'
    r_bcf.font.size = Pt(8.5)
    r_bcf.font.bold = True

    # Rows BC
    df_bc = df_uc[df_uc['Tipo'] == 'COMPLEMENTAR'].reset_index(drop=True)
    for i, bc_ref in enumerate(uc['bc']):
        row_bc = table.add_row()
        c_ref = row_bc.cells[0].merge(row_bc.cells[1])
        c_qtd = row_bc.cells[2]
        c_fv = row_bc.cells[3]
        
        set_cell_margins(c_ref, top=60, bottom=60, left=80, right=80)
        set_cell_margins(c_qtd, top=60, bottom=60, left=40, right=40)
        set_cell_margins(c_fv, top=60, bottom=60, left=40, right=40)
        
        row_data = df_bc.iloc[i].to_dict() if i < len(df_bc) else {}
        qtd_str, tipo_str = determine_availability_and_type(row_data, bc_ref)
        
        add_abnt_reference_paragraph(c_ref, ref_idx, bc_ref)
        
        p_q = c_qtd.paragraphs[0]
        p_q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_q.paragraph_format.space_after = Pt(0)
        r_q = p_q.add_run(qtd_str)
        r_q.font.name = 'Arial'
        r_q.font.size = Pt(9.5)
        if "0" in qtd_str:
            r_q.font.color.rgb = RGBColor(185, 28, 28)
            r_q.font.bold = True
        
        p_f = c_fv.paragraphs[0]
        p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_f.paragraph_format.space_after = Pt(0)
        r_f = p_f.add_run(tipo_str)
        r_f.font.name = 'Arial'
        r_f.font.size = Pt(9.5)
        r_f.font.bold = True
        
        ref_idx += 1

    # ROW Parecer do Docente (span 4)
    row_par = table.add_row()
    c_par = row_par.cells[0].merge(row_par.cells[1]).merge(row_par.cells[2]).merge(row_par.cells[3])
    set_cell_margins(c_par, top=100, bottom=120, left=120, right=120)
    p_par = c_par.paragraphs[0]
    p_par.paragraph_format.space_after = Pt(4)
    p_par.paragraph_format.line_spacing = 1.15
    r_parl = p_par.add_run("Parecer do docente:\n")
    r_parl.font.name = 'Arial'
    r_parl.font.size = Pt(10)
    r_parl.font.bold = True
    
    r_part = p_par.add_run(
        "(Justifique sempre que precisar substituir ou atualizar em função da necessidade de aquisição "
        "de novos títulos, ou pela ausência no acervo. Destacar os pontos principais que justificam a adequação "
        "de cada título à ementa e aos objetivos da UC. Discutir se a quantidade e a qualidade dos materiais previstos "
        "atendem às necessidades da UC. Propor substituição, caso necessário).\n\n\n\n"
    )
    r_part.font.name = 'Arial'
    r_part.font.size = Pt(9)
    r_part.font.italic = True
    r_part.font.color.rgb = RGBColor(71, 85, 105)

    # ROW Cabeçalho Assinaturas (span 4)
    row_ass_head = table.add_row()
    c_ass_h = row_ass_head.cells[0].merge(row_ass_head.cells[1]).merge(row_ass_head.cells[2]).merge(row_ass_head.cells[3])
    set_cell_background(c_ass_h, HEX_GRAY_HEADER)
    set_cell_margins(c_ass_h, top=80, bottom=80, left=100, right=100)
    p_ass_h = c_ass_h.paragraphs[0]
    p_ass_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ass_h.paragraph_format.space_after = Pt(0)
    r_ass_h = p_ass_h.add_run("Assinaturas validando as bibliografias pelos docentes e referendado pelo NDE")
    r_ass_h.font.name = 'Arial'
    r_ass_h.font.size = Pt(10.5)
    r_ass_h.font.bold = True

    # ROW Docente
    row_doc = table.add_row()
    c_doc_l = row_doc.cells[0]
    c_doc_v = row_doc.cells[1].merge(row_doc.cells[2]).merge(row_doc.cells[3])
    set_cell_margins(c_doc_l, top=70, bottom=70, left=100, right=100)
    set_cell_margins(c_doc_v, top=70, bottom=70, left=100, right=100)
    
    p_dl = c_doc_l.paragraphs[0]
    p_dl.paragraph_format.space_after = Pt(0)
    r_dl = p_dl.add_run("Docente:")
    r_dl.font.name = 'Arial'
    r_dl.font.size = Pt(10)
    r_dl.font.bold = True
    
    p_dv = c_doc_v.paragraphs[0]
    p_dv.paragraph_format.space_after = Pt(0)

    # ROWS Membros NDE (1 a 8)
    for m in NDE_MEMBERS:
        row_m = table.add_row()
        c_ml = row_m.cells[0]
        c_mv = row_m.cells[1].merge(row_m.cells[2]).merge(row_m.cells[3])
        set_cell_margins(c_ml, top=60, bottom=60, left=100, right=100)
        set_cell_margins(c_mv, top=60, bottom=60, left=100, right=100)
        
        p_ml = c_ml.paragraphs[0]
        p_ml.paragraph_format.space_after = Pt(0)
        r_ml = p_ml.add_run(m)
        r_ml.font.name = 'Arial'
        r_ml.font.size = Pt(9.5)
        
        p_mv = c_mv.paragraphs[0]
        p_mv.paragraph_format.space_after = Pt(0)

    # Definir larguras de colunas da tabela e aplicar bordas em TODAS as células
    col_widths = [Inches(3.35), Inches(1.46), Inches(1.00), Inches(0.89)]
    for row in table.rows:
        # Previne quebra de linha no meio da página no Google Docs
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        
        for j, cell in enumerate(row.cells):
            set_cell_borders(cell, color=HEX_BORDER, sz="6", val="single")
            if j < len(col_widths):
                try:
                    cell.width = col_widths[j]
                except:
                    pass

def generate_relatorio_nde():
    print(f"--- Iniciando Geração do Relatório NDE Gestão Ambiental (Google Docs Ready) ---")
    print(f"Destino: {OUTPUT_DOCX_PATH}")
    
    ucs_info, df_all = load_all_data()
    print(f"Carregadas {len(ucs_info)} UCs e {len(df_all)} registros de auditoria.")
    
    doc = Document()
    
    # Configurar margens A4 idênticas ao modelo
    for s in doc.sections:
        s.page_width = Inches(8.27)   # A4 210mm
        s.page_height = Inches(11.69) # A4 297mm
        s.top_margin = Inches(0.40)   # 27 pt
        s.bottom_margin = Inches(0.90)# 66 pt
        s.left_margin = Inches(0.69)  # 49.6 pt
        s.right_margin = Inches(0.69) # 49.6 pt
        
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = COLOR_DARK
    
    # Capa / Apresentação Institucional do NDE
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_head.paragraph_format.space_after = Pt(4)
    r_h1 = p_head.add_run("INSTITUTO FEDERAL DE SANTA CATARINA — CÂMPUS GAROPABA\n")
    r_h1.font.name = 'Arial'
    r_h1.font.bold = True
    r_h1.font.size = Pt(11)
    r_h1.font.color.rgb = COLOR_IFSC_GREEN
    
    r_h2 = p_head.add_run("NÚCLEO DOCENTE ESTRUTURANTE (NDE) — CST EM GESTÃO AMBIENTAL\n")
    r_h2.font.name = 'Arial'
    r_h2.font.bold = True
    r_h2.font.size = Pt(12.5)
    r_h2.font.color.rgb = COLOR_DARK
    
    r_h3 = p_head.add_run("Relatório Consolidado de Adequação e Validação de Bibliografias (PPC 2026)\n")
    r_h3.font.name = 'Arial'
    r_h3.font.bold = True
    r_h3.font.size = Pt(10.5)
    r_h3.font.color.rgb = RGBColor(100, 116, 139)
    
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(12)
    p_intro.paragraph_format.line_spacing = 1.15
    r_int = p_intro.add_run(
        "Este documento consolida as fichas individuais de análise da adequação bibliográfica das "
        "33 Unidades Curriculares do Curso Superior de Tecnologia em Gestão Ambiental, com o cruzamento oficial "
        "entre o PPC 2026 e o catálogo de acervo e exemplares do Câmpus Garopaba (Sistema Sophia / Biblioteca Virtual). "
        "O formulário segue rigorosamente o padrão normativo institucional para conferência pelo docente e referendo pelo NDE."
    )
    r_int.font.name = 'Arial'
    r_int.font.size = Pt(9.5)
    
    # Adicionar fichas de cada UC
    for idx, uc in enumerate(ucs_info):
        print(f"[{idx+1}/33] Gerando Ficha NDE para: {uc['uc_nome']}")
        doc.add_page_break()
        df_uc = df_all[df_all['UC'].apply(lambda x: normalize_str(x) == normalize_str(uc['uc_nome']))]
        build_uc_nde_table(doc, uc, df_uc)
        
    doc.save(OUTPUT_DOCX_PATH)
    print(f"\nDocumento NDE gerado com sucesso em:\n{OUTPUT_DOCX_PATH}")
    file_size = os.path.getsize(OUTPUT_DOCX_PATH)
    print(f"Tamanho do arquivo: {file_size / 1024:.1f} KB")

if __name__ == '__main__':
    generate_relatorio_nde()
