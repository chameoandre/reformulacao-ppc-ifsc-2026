#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Gerador Fiel do Ementário do PPC CST Gestão Ambiental em Formato Word (.docx) com Sumário Clicável e Navegação Rápida:
Compatível 100% com Microsoft Word, LibreOffice e Google Docs (incluindo painel lateral de estrutura).

Recorte literal da Seção de Ementas do PPC CST em Gestão Ambiental (IFSC Garopaba).
"""

import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

BASE_DIR = "/Users/chameoandre/Google-Drive-chameoandre/INSTITUTO-FEDERAL-SANTA-CATARINA/CURSOS/TECNICO/informatica-integrado/reformulacao-ppc-informatica-administracao-integrado/superior-gestao-ambiental"
CHECK_DIR = os.path.join(BASE_DIR, "revisao-equipe", "checagem-biblioteca")
TXT_PATH = os.path.join(CHECK_DIR, "Unidades Curriculares Gestão Ambiental.txt")
DOCX_OUTPUT_PATH = os.path.join(CHECK_DIR, "Ementario_Completo_PPC_Gestao_Ambiental_Revisao_Biblioteca.docx")

# Colors matching IFSC PPC
COLOR_IFSC_GREEN = RGBColor(16, 140, 80)      # #108C50
COLOR_DARK = RGBColor(15, 23, 42)             # #0F172A
COLOR_MUTED = RGBColor(100, 116, 139)         # #64748B
COLOR_BLUE = RGBColor(2, 132, 199)            # #0284C7
HEX_IFSC_GREEN = "108C50"
HEX_LIGHT_GREEN = "EBF8F0"
HEX_LIGHT_GRAY = "F1F5F9"
HEX_BORDER = "CBD5E1"
HEX_BLUE = "0284C7"

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_bookmark(paragraph, bookmark_id, bookmark_name):
    p = paragraph._p
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bookmark_id))
    bm_start.set(qn('w:name'), bookmark_name)
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bookmark_id))
    p.append(bm_start)
    p.append(bm_end)

def add_internal_hyperlink(paragraph, anchor_name, text, color=HEX_BLUE, bold=True, size=9.0):
    p = paragraph._p
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor_name)
    hyperlink.set(qn('w:history'), '1')
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
        
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
        
    if size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size * 2)))
        rPr.append(sz)
        
    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    
    hyperlink.append(new_run)
    p.append(hyperlink)

try:
    import analise_acervo_biblioteca_ga as ga
except ImportError:
    ga = None

def get_bold_title_runs(ref_raw):
    raw = ref_raw.strip()
    if any(k in raw for k in ["LIVRO didático", "Material didático", "Fundo Nacional"]):
        return [(raw, False)]
        
    meta = ga.extract_metadata(raw) if ga else {}
    titulo = meta.get("titulo", "").strip()
    titulo_curto = meta.get("titulo_curto", "").strip()
    
    candidates = [titulo, titulo_curto]
    for cand in candidates:
        if not cand: continue
        idx = raw.find(cand)
        if idx != -1:
            prefix = raw[:idx]
            title_text = raw[idx:idx+len(cand)]
            suffix = raw[idx+len(cand):]
            runs = []
            if prefix: runs.append((prefix, False))
            if title_text: runs.append((title_text, True))
            if suffix: runs.append((suffix, False))
            return runs
            
        pattern = re.escape(cand)
        m = re.search(pattern, raw, re.IGNORECASE)
        if m:
            prefix = raw[:m.start()]
            title_text = raw[m.start():m.end()]
            suffix = raw[m.end():]
            runs = []
            if prefix: runs.append((prefix, False))
            if title_text: runs.append((title_text, True))
            if suffix: runs.append((suffix, False))
            return runs
            
    m = re.match(r'^([A-ZÁÉÍÓÚÂÊÔÃÕÇ\s\.,;&\(\)\-]+?\.\s+)(.+?)(?=\.\s+(?:[A-Z][a-z]+|\d+\.\s*ed|[A-ZÁÉÍÓÚÂÊÔÃÕÇ\s]+:|$))(.*)$', raw)
    if m:
        return [(m.group(1), False), (m.group(2), True), (m.group(3), False)]
        
    return [(raw, False)]

def add_abnt_reference_paragraph(cell, ref_text):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    raw = ref_text.strip()
    if not raw:
        return p
    runs_data = get_bold_title_runs(raw)
    for text_part, is_bold in runs_data:
        r = p.add_run(text_part)
        r.font.size = Pt(9)
        if is_bold:
            r.font.bold = True
    return p

def clean_val(t):
    if not t: return ''
    return t.replace('**', '').replace('*', '').strip()

def merge_reference_lines(raw_text):
    # 1. Split concatenated refs on same line
    t = raw_text
    t = re.sub(r'(\b\d{4}\.)\s+([A-ZÁÉÍÓÚÂÊÔÃÕÇ]{2,}(?:\s+[A-Z]\.|\s+[A-ZÁÉÍÓÚÂÊÔÃÕÇ]+)*[,\s]\s*)', r'\1\n\2', t)
    t = re.sub(r'(\b\d{4}\.)\s+(______[\.\s])', r'\1\n\2', t)
    t = re.sub(r'(\b______\.\s+[^\n]+?\b\d{4}\.)\s+([A-ZÁÉÍÓÚÂÊÔÃÕÇ]{2,},\s+|______)', r'\1\n\2', t)

    lines = [l.strip() for l in t.split('\n') if l.strip()]
    merged = []
    for l in lines:
        if l.startswith('(*)') or 'CH Total' in l or 'CH EaD' in l:
            continue
        if not merged:
            merged.append(l)
        else:
            is_new_ref = False
            if not (merged[-1].rstrip().endswith(';') or merged[-1].rstrip().endswith('&')):
                if re.match(r'^[A-ZÁÉÍÓÚÂÊÔÃÕÇ]{2,}(?:\s+[A-ZÁÉÍÓÚÂÊÔÃÕÇ]{2,})*,\s+[A-Za-zÀ-ÿ]', l):
                    is_new_ref = True
                elif re.match(r'^[A-ZÁÉÍÓÚÂÊÔÃÕÇ]{2,}\s+[A-Z]\.\s*(?:[A-Z]\.)?\s+[A-Za-zÀ-ÿ]', l):
                    is_new_ref = True
                elif re.match(r'^(BRASIL|EMBRAPA|IBGE|MMA|MEC|UNESCO|WHO|ONU|CONAMA|BANCO MUNDIAL|INSTITUTO|MINISTÉRIO|SECRETARIA|FNDE|AGÊNCIA)\b', l, re.IGNORECASE):
                    is_new_ref = True
                elif re.match(r'^\d+[\.\)]\s*[A-ZÁÉÍÓÚÂÊÔÃÕÇ]{2,}', l):
                    is_new_ref = True
                elif l.startswith('LIVRO didático') or l.startswith('Material didático'):
                    is_new_ref = True
                elif l.startswith('______') or l.startswith('____.'):
                    is_new_ref = True
                
            if is_new_ref:
                merged.append(l)
            else:
                merged[-1] = merged[-1] + ' ' + l
                
    return [r.strip() for r in merged if len(r.strip()) > 10]

def parse_txt_ementas():
    with open(TXT_PATH, "r", encoding="utf-8-sig") as f:
        text = f.read()

    raw_blocks = text.split('------------------------------------------------------------')
    if len(raw_blocks) < 10:
        raw_blocks = re.split(r'\n(?=[A-ZÁÉÍÓÚÂÊÔÃÕÇ\s\-\/\(\)]{3,60}\n\s*Semestre\s*:)', text)
        if len(raw_blocks) < 10:
            raw_blocks = [u.strip() for u in text.split('Unidade Curricular:') if u.strip()]

    ucs_info = []
    idx = 0
    for block in raw_blocks:
        if 'Semestre:' not in block:
            continue
        idx += 1
        
        # UC Name
        b_lines = [l.strip() for l in block.strip().split('\n') if l.strip()]
        uc_nome = "UC DESCONHECIDA"
        for i, l in enumerate(b_lines):
            if 'Semestre:' in l:
                if i > 0:
                    uc_nome = b_lines[i-1]
                break
                
        # Clean uc_nome
        uc_nome = re.sub(r'^\d+[\.\-\s]+', '', uc_nome).strip()
        uc_nome = uc_nome.replace('UC:', '').replace('Unidade Curricular:', '').strip()
        
        # Semestre
        sem_match = re.search(r'Semestre\s*:\s*(\d+)', block)
        sem_val = sem_match.group(1) if sem_match else "1"
        
        # Cargas Horárias
        ch_match = re.search(r'CH Total\*?:\s*([^\n]+)', block)
        ch_val = ch_match.group(1).strip() if ch_match else "40 h"
        
        ch_ead_match = re.search(r'CH EaD\*?:\s*([^\n]+)', block)
        ch_ead_val = ch_ead_match.group(1).strip() if ch_ead_match else "00 h"
        
        ch_ext_match = re.search(r'CH Extens[aã]o:\s*([^\n]+)', block)
        ch_ext_val = ch_ext_match.group(1).strip() if ch_ext_match else "00 h"

        ch_pres_match = re.search(r'CH Presencial:\s*([^\n]+)', block)
        ch_pres_val = ch_pres_match.group(1).strip() if ch_pres_match else ""

        # Objectives
        obj_items = []
        if 'Objetivos:' in block:
            obj_part = block.split('Objetivos:')[1]
            if 'Conteúdos:' in obj_part:
                obj_part = obj_part.split('Conteúdos:')[0]
            for l in obj_part.split('\n'):
                l_str = l.strip()
                if l_str.startswith('*') or l_str.startswith('•') or l_str.startswith('-'):
                    obj_items.append(l_str.lstrip('*•- ').strip())
                elif l_str and len(l_str) > 15:
                    obj_items.append(l_str)

        # Contents
        cont_text = ""
        if 'Conteúdos:' in block:
            cont_part = block.split('Conteúdos:')[1]
            if 'Estratégias de ensino' in cont_part:
                cont_part = cont_part.split('Estratégias de ensino')[0]
            cont_text = cont_part.strip()

        # Strategies
        estr_text = ""
        if 'Estratégias de ensino' in block:
            estr_part = block.split('Estratégias de ensino')[1]
            if 'Bibliografia Básica:' in estr_part:
                estr_part = estr_part.split('Bibliografia Básica:')[0]
            estr_text = estr_part.strip().lstrip(': \n')

        # BB
        bb_list = []
        if 'Bibliografia Básica:' in block:
            bb_part = block.split('Bibliografia Básica:')[1]
            if 'Bibliografia Complementar:' in bb_part:
                bb_part = bb_part.split('Bibliografia Complementar:')[0]
            else:
                bb_part = bb_part.split('(*)')[0]
            bb_list = merge_reference_lines(bb_part)
                    
        # BC
        bc_list = []
        if 'Bibliografia Complementar:' in block:
            bc_part = block.split('Bibliografia Complementar:')[1].split('(*)')[0]
            bc_list = merge_reference_lines(bc_part)

        ucs_info.append({
            'id': idx,
            'uc_nome': uc_nome,
            'semestre': sem_val,
            'ch_total': ch_val,
            'ch_ead': ch_ead_val,
            'ch_ext': ch_ext_val,
            'ch_pres': ch_pres_val,
            'objetivos': obj_items,
            'conteudos': cont_text,
            'estrategias': estr_text,
            'qtd_bb': len(bb_list),
            'qtd_bc': len(bc_list),
            'bb': bb_list,
            'bc': bc_list
        })

    return ucs_info

def generate_ga_docx():
    print(f"Gerando Documento Word (.docx) Gestão Ambiental em:\n{DOCX_OUTPUT_PATH}")
    uc_tables = parse_txt_ementas()
    
    doc = Document()
    
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(9.5)
    style_normal.font.color.rgb = COLOR_DARK

    # Cabeçalho Oficial
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_h1 = p_head.add_run("INSTITUTO FEDERAL DE SANTA CATARINA — CÂMPUS GAROPABA\n")
    r_h1.font.bold = True
    r_h1.font.size = Pt(11)
    r_h1.font.color.rgb = COLOR_IFSC_GREEN
    
    r_h2 = p_head.add_run("PROJETO PEDAGÓGICO DE CURSO (PPC 2026)\nCURSO SUPERIOR DE TECNOLOGIA EM GESTÃO AMBIENTAL\n")
    r_h2.font.bold = True
    r_h2.font.size = Pt(13.5)
    r_h2.font.color.rgb = COLOR_DARK
    
    r_h3 = p_head.add_run("Seção da Estrutura Curricular & Ementário das 33 Unidades Curriculares\n")
    r_h3.font.bold = True
    r_h3.font.size = Pt(10.5)
    r_h3.font.color.rgb = COLOR_MUTED

    # Sumário Bookmark
    p_sum_title = doc.add_paragraph()
    p_sum_title.paragraph_format.space_before = Pt(8)
    p_sum_title.paragraph_format.space_after = Pt(4)
    add_bookmark(p_sum_title, 100, "sumario_geral")
    
    r_st = p_sum_title.add_run("📑 SUMÁRIO GERAL CLICÁVEL — NAVEGAÇÃO RÁPIDA PELAS 33 EMENTAS")
    r_st.font.bold = True
    r_st.font.size = Pt(11)
    r_st.font.color.rgb = COLOR_IFSC_GREEN

    p_sum_desc = doc.add_paragraph()
    p_sum_desc.paragraph_format.space_after = Pt(8)
    r_sd = p_sum_desc.add_run("Clique sobre qualquer Unidade Curricular abaixo para navegar instantaneamente até a sua ementa e bibliografia completa:")
    r_sd.font.italic = True
    r_sd.font.size = Pt(8.5)
    r_sd.font.color.rgb = COLOR_MUTED

    # Group UCs by Semesters: Sem 1-2, Sem 3-4, Sem 5
    tbl_toc = doc.add_table(rows=14, cols=3)
    tbl_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_toc, color=HEX_BORDER)
    
    col_headers_toc = ["1º e 2º SEMESTRES", "3º e 4º SEMESTRES", "5º SEMESTRE"]
    w_toc_col = Inches(2.26)
    
    for c_i in range(3):
        c = tbl_toc.cell(0, c_i)
        c.width = w_toc_col
        set_cell_background(c, HEX_LIGHT_GREEN)
        set_cell_margins(c, top=70, bottom=70, left=80, right=80)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(col_headers_toc[c_i])
        r.font.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = COLOR_IFSC_GREEN

    # Sem 1 & 2: UCs 1 a 13
    # Sem 3 & 4: UCs 14 a 26
    # Sem 5: UCs 27 a 33
    for r_idx in range(1, 14):
        # Col 0 (UCs 1-13)
        uc_idx_1 = r_idx
        c0 = tbl_toc.cell(r_idx, 0)
        c0.width = w_toc_col
        set_cell_margins(c0, top=35, bottom=35, left=60, right=60)
        if uc_idx_1 <= 13 and uc_idx_1 <= len(uc_tables):
            u_item = uc_tables[uc_idx_1 - 1]
            p0 = c0.paragraphs[0]
            add_internal_hyperlink(p0, f"uc_{uc_idx_1}", f"{uc_idx_1:02d}. {u_item['uc_nome']} (Sem {u_item['semestre']})", color=HEX_BLUE, bold=False, size=8.0)

        # Col 1 (UCs 14-26)
        uc_idx_2 = 13 + r_idx
        c1 = tbl_toc.cell(r_idx, 1)
        c1.width = w_toc_col
        set_cell_margins(c1, top=35, bottom=35, left=60, right=60)
        if uc_idx_2 <= 26 and uc_idx_2 <= len(uc_tables):
            u_item = uc_tables[uc_idx_2 - 1]
            p1 = c1.paragraphs[0]
            add_internal_hyperlink(p1, f"uc_{uc_idx_2}", f"{uc_idx_2:02d}. {u_item['uc_nome']} (Sem {u_item['semestre']})", color=HEX_BLUE, bold=False, size=8.0)

        # Col 2 (UCs 27-33)
        uc_idx_3 = 26 + r_idx
        c2 = tbl_toc.cell(r_idx, 2)
        c2.width = w_toc_col
        set_cell_margins(c2, top=35, bottom=35, left=60, right=60)
        if uc_idx_3 <= 33 and uc_idx_3 <= len(uc_tables):
            u_item = uc_tables[uc_idx_3 - 1]
            p2 = c2.paragraphs[0]
            add_internal_hyperlink(p2, f"uc_{uc_idx_3}", f"{uc_idx_3:02d}. {u_item['uc_nome']} (Sem {u_item['semestre']})", color=HEX_BLUE, bold=False, size=8.0)

    doc.add_page_break()

    # Seção do Ementário
    p_sec2 = doc.add_paragraph()
    r_s2 = p_sec2.add_run("Ementário das 33 Unidades Curriculares — CST em Gestão Ambiental")
    r_s2.font.bold = True
    r_s2.font.size = Pt(12.5)
    r_s2.font.color.rgb = COLOR_IFSC_GREEN
    p_sec2.paragraph_format.space_after = Pt(12)

    for idx, uc in enumerate(uc_tables, 1):
        p_uc_title = doc.add_paragraph()
        p_uc_title.paragraph_format.space_before = Pt(6)
        p_uc_title.paragraph_format.space_after = Pt(3)
        add_bookmark(p_uc_title, 200 + idx, f"uc_{idx}")
        
        r_t_h = p_uc_title.add_run(f"UC {idx:02d} — {uc['uc_nome']}")
        r_t_h.font.bold = True
        r_t_h.font.size = Pt(11)
        r_t_h.font.color.rgb = COLOR_IFSC_GREEN
        
        tbl_uc = doc.add_table(rows=8, cols=3)
        tbl_uc.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl_uc, color=HEX_BORDER)
        
        w_col0 = Inches(4.2)
        w_col1 = Inches(1.3)
        w_col2 = Inches(1.3)
        
        # Row 0: Header
        c_uc = tbl_uc.cell(0, 0)
        c_uc.width = w_col0
        set_cell_background(c_uc, HEX_LIGHT_GREEN)
        set_cell_margins(c_uc, top=70, bottom=70, left=90, right=90)
        p_u = c_uc.paragraphs[0]
        r_u_tag = p_u.add_run("Unidade Curricular:\n")
        r_u_tag.font.bold = True
        r_u_tag.font.size = Pt(8.5)
        r_u_tag.font.color.rgb = COLOR_IFSC_GREEN
        r_u_name = p_u.add_run(uc["uc_nome"])
        r_u_name.font.bold = True
        r_u_name.font.size = Pt(11)
        r_u_name.font.color.rgb = COLOR_DARK
        
        c_sem = tbl_uc.cell(0, 1)
        c_sem.merge(tbl_uc.cell(0, 2))
        set_cell_margins(c_sem, top=70, bottom=70, left=90, right=90)
        p_s = c_sem.paragraphs[0]
        r_s_tag = p_s.add_run("Semestre: ")
        r_s_tag.font.bold = True
        r_s_tag.font.size = Pt(8.5)
        r_s_tag.font.color.rgb = COLOR_IFSC_GREEN
        r_s_val = p_s.add_run(str(uc["semestre"]))
        r_s_val.font.bold = True
        r_s_val.font.size = Pt(9.5)

        # Row 1: CH EaD e CH Total
        c_r1_0 = tbl_uc.cell(1, 0)
        set_cell_margins(c_r1_0, top=40, bottom=40, left=90, right=90)
        p_r1_0 = c_r1_0.paragraphs[0]
        p_r1_0.add_run("Matriz Curricular CST Gestão Ambiental").font.size = Pt(8)
        p_r1_0.runs[0].font.color.rgb = COLOR_MUTED
        
        c_ead = tbl_uc.cell(1, 1)
        set_cell_margins(c_ead, top=40, bottom=40, left=90, right=90)
        p_ead = c_ead.paragraphs[0]
        p_ead.add_run("CH EaD: ").font.bold = True
        p_ead.runs[0].font.color.rgb = COLOR_IFSC_GREEN
        p_ead.runs[0].font.size = Pt(8.5)
        p_ead.add_run(uc["ch_ead"]).font.bold = True
        p_ead.runs[1].font.size = Pt(9)
        
        c_tot = tbl_uc.cell(1, 2)
        set_cell_margins(c_tot, top=40, bottom=40, left=90, right=90)
        p_tot = c_tot.paragraphs[0]
        p_tot.add_run("CH Total*: ").font.bold = True
        p_tot.runs[0].font.color.rgb = COLOR_IFSC_GREEN
        p_tot.runs[0].font.size = Pt(8.5)
        p_tot.add_run(uc["ch_total"]).font.bold = True
        p_tot.runs[1].font.size = Pt(9)

        # Row 2: Objetivos
        c_obj = tbl_uc.cell(2, 0)
        c_obj.merge(tbl_uc.cell(2, 2))
        set_cell_margins(c_obj, top=70, bottom=70, left=90, right=90)
        p_obj_h = c_obj.paragraphs[0]
        r_oh = p_obj_h.add_run("Objetivos:\n")
        r_oh.font.bold = True
        r_oh.font.size = Pt(9)
        r_oh.font.color.rgb = COLOR_IFSC_GREEN
        
        for obj_it in uc["objetivos"]:
            p_it = c_obj.add_paragraph()
            p_it.paragraph_format.left_indent = Inches(0.2)
            p_it.paragraph_format.space_after = Pt(2)
            p_it.add_run("• " + obj_it).font.size = Pt(9)

        # Row 3: Conteúdos
        c_cont = tbl_uc.cell(3, 0)
        c_cont.merge(tbl_uc.cell(3, 2))
        set_cell_margins(c_cont, top=70, bottom=70, left=90, right=90)
        p_cont_h = c_cont.paragraphs[0]
        r_ch = p_cont_h.add_run("Conteúdos:\n")
        r_ch.font.bold = True
        r_ch.font.size = Pt(9)
        r_ch.font.color.rgb = COLOR_IFSC_GREEN
        
        for cont_line in uc["conteudos"].split('\n'):
            if cont_line.strip():
                p_c = c_cont.add_paragraph()
                p_c.paragraph_format.space_after = Pt(3)
                p_c.add_run(cont_line.strip()).font.size = Pt(9)

        # Row 4: Estratégias
        c_estr = tbl_uc.cell(4, 0)
        c_estr.merge(tbl_uc.cell(4, 2))
        set_cell_margins(c_estr, top=70, bottom=70, left=90, right=90)
        p_estr_h = c_estr.paragraphs[0]
        r_eh = p_estr_h.add_run("Estratégias de Ensino e Aprendizagem:\n")
        r_eh.font.bold = True
        r_eh.font.size = Pt(9)
        r_eh.font.color.rgb = COLOR_IFSC_GREEN
        
        for estr_line in uc["estrategias"].split('\n'):
            if estr_line.strip():
                p_e = c_estr.add_paragraph()
                p_e.paragraph_format.space_after = Pt(3)
                p_e.add_run(estr_line.strip()).font.size = Pt(9)

        # Row 5: Bibliografia Básica
        c_bb = tbl_uc.cell(5, 0)
        c_bb.merge(tbl_uc.cell(5, 2))
        set_cell_margins(c_bb, top=70, bottom=70, left=90, right=90)
        set_cell_background(c_bb, HEX_LIGHT_GREEN)
        p_bb_h = c_bb.paragraphs[0]
        r_bbh = p_bb_h.add_run("Bibliografia Básica:\n")
        r_bbh.font.bold = True
        r_bbh.font.size = Pt(9)
        r_bbh.font.color.rgb = COLOR_IFSC_GREEN
        
        for bb_item in uc["bb"]:
            if bb_item.strip():
                add_abnt_reference_paragraph(c_bb, bb_item)

        # Row 6: Bibliografia Complementar
        c_bc = tbl_uc.cell(6, 0)
        c_bc.merge(tbl_uc.cell(6, 2))
        set_cell_margins(c_bc, top=70, bottom=70, left=90, right=90)
        p_bc_h = c_bc.paragraphs[0]
        r_bch = p_bc_h.add_run("Bibliografia Complementar:\n")
        r_bch.font.bold = True
        r_bch.font.size = Pt(9)
        r_bch.font.color.rgb = COLOR_IFSC_GREEN
        
        for bc_item in uc["bc"]:
            if bc_item.strip():
                add_abnt_reference_paragraph(c_bc, bc_item)
                
        # Row 7: Navegação de Retorno ao Sumário
        c_r7 = tbl_uc.cell(7, 0)
        c_r7.merge(tbl_uc.cell(7, 2))
        set_cell_margins(c_r7, top=30, bottom=30, left=90, right=90)
        set_cell_background(c_r7, HEX_LIGHT_GRAY)
        p_end = c_r7.paragraphs[0]
        p_end.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_internal_hyperlink(p_end, "sumario_geral", "🔝 Voltar ao Sumário Geral", color=HEX_BLUE, bold=True, size=8.0)
        p_end.add_run(f"   |   PPC CST Gestão Ambiental (IFSC Garopaba) • UC {idx:02d}").font.size = Pt(7.5)
        p_end.runs[-1].font.color.rgb = COLOR_MUTED

        if idx < len(uc_tables):
            doc.add_page_break()

    doc.save(DOCX_OUTPUT_PATH)
    print(f"Documento Word (.docx) Gestão Ambiental salvo com sucesso em:\n{DOCX_OUTPUT_PATH}")

if __name__ == "__main__":
    generate_ga_docx()
